import os
import numpy as np
import pandas as pd
import tensorflow as tf
from tensorflow.keras.models import Model
from tensorflow.keras.layers import Input, Dense, LayerNormalization, Dropout, GlobalAveragePooling1D
from tensorflow.keras.optimizers import Adam
from tensorflow.keras.utils import to_categorical
from sklearn.model_selection import train_test_split
from sklearn.preprocessing import LabelEncoder
import cv2

# Paths to your dataset
dataset_path = r'C:\Users\BEST TECH\Downloads\SUE+UKH - Copy'
csv_path = r'C:\Users\BEST TECH\Downloads\UKH+SUE-labels.csv'

# Image size for Vision Transformer
IMG_SIZE = 128
PATCH_SIZE = 16
NUM_PATCHES = (IMG_SIZE // PATCH_SIZE) ** 2

# Load and preprocess dataset
def load_and_preprocess_dataset():
    labels_df = pd.read_csv(csv_path)

    # Get valid folder IDs
    valid_sentence_ids = [folder_name for folder_name in os.listdir(dataset_path) if
                          os.path.isdir(os.path.join(dataset_path, folder_name))]
    filtered_labels_df = labels_df[labels_df['ID'].isin(valid_sentence_ids)].copy()

    # Encode labels
    label_encoder = LabelEncoder()
    filtered_labels_df['label'] = label_encoder.fit_transform(filtered_labels_df['sentence'])

    images = []
    labels = []

    for sentence_id in os.listdir(dataset_path):
        sentence_folder = os.path.join(dataset_path, sentence_id)
        if os.path.isdir(sentence_folder) and sentence_id in filtered_labels_df['ID'].values:
            label = filtered_labels_df[filtered_labels_df['ID'] == sentence_id]['label'].values[0]
            for img_file in os.listdir(sentence_folder):
                img_path = os.path.join(sentence_folder, img_file)

                # Check for valid image file
                img = cv2.imread(img_path, cv2.IMREAD_GRAYSCALE)
                if img is None:
                    print(f"Skipping unreadable image: {img_path}")
                    continue

                # Resize and normalize
                img = cv2.resize(img, (IMG_SIZE, IMG_SIZE)) / 255.0
                images.append(img)
                labels.append(label)

    images = np.array(images, dtype=np.float32).reshape(-1, IMG_SIZE, IMG_SIZE, 1)
    labels = np.array(labels)
    print(f"Loaded {len(images)} images with {len(labels)} labels.")
    return images, labels, label_encoder

images, labels, label_encoder = load_and_preprocess_dataset()

# Train-test split
X_train, X_temp, y_train, y_temp = train_test_split(images, labels, test_size=0.3, random_state=42)
X_val, X_test, y_val, y_test = train_test_split(X_temp, y_temp, test_size=0.5, random_state=42)

# Convert labels to categorical
num_classes = len(np.unique(labels))
y_train = to_categorical(y_train, num_classes)
y_val = to_categorical(y_val, num_classes)
y_test = to_categorical(y_test, num_classes)

# Helper function: Patch Embedding Layer
def create_patches(images, patch_size):
    batch_size = tf.shape(images)[0]
    patches = tf.image.extract_patches(
        images=images,
        sizes=[1, patch_size, patch_size, 1],
        strides=[1, patch_size, patch_size, 1],
        rates=[1, 1, 1, 1],
        padding='VALID'
    )
    patch_dim = patches.shape[-1]
    patches = tf.reshape(patches, [batch_size, -1, patch_dim])
    return patches

# Vision Transformer (ViT) Model
def build_vit(input_shape, num_classes, num_patches, embed_dim=128, num_heads=8, ff_dim=256, num_layers=4):
    inputs = Input(shape=input_shape)

    # Create patches
    patches = tf.keras.layers.Lambda(create_patches, arguments={'patch_size': PATCH_SIZE})(inputs)

    # Patch embedding
    patch_embed = Dense(embed_dim)(patches)

    # Add positional embeddings
    positions = tf.range(start=0, limit=num_patches, delta=1)
    pos_embed = tf.keras.layers.Embedding(input_dim=num_patches, output_dim=embed_dim)(positions)
    patch_embed += pos_embed

    # Transformer layers
    for _ in range(num_layers):
        # Multi-head self-attention
        attention_output = tf.keras.layers.MultiHeadAttention(num_heads=num_heads, key_dim=embed_dim)(patch_embed, patch_embed)
        attention_output = Dropout(0.1)(attention_output)
        attention_output = LayerNormalization(epsilon=1e-6)(attention_output + patch_embed)

        # Feed-forward network
        ff_output = Dense(ff_dim, activation='relu')(attention_output)
        ff_output = Dense(embed_dim)(ff_output)
        ff_output = Dropout(0.1)(ff_output)
        patch_embed = LayerNormalization(epsilon=1e-6)(ff_output + attention_output)

    # Classification head
    x = GlobalAveragePooling1D()(patch_embed)
    x = Dropout(0.5)(x)
    outputs = Dense(num_classes, activation='softmax')(x)

    model = Model(inputs, outputs)
    return model

# Build and compile the model
input_shape = (IMG_SIZE, IMG_SIZE, 1)
vit_model = build_vit(input_shape, num_classes, num_patches=NUM_PATCHES)
vit_model.compile(optimizer=Adam(learning_rate=1e-4), loss='categorical_crossentropy', metrics=['accuracy'])
vit_model.summary()

# Train the model
history = vit_model.fit(X_train, y_train, epochs=20, batch_size=32, validation_data=(X_val, y_val))

# Save the model
vit_model.save('vit_kurdish_handwriting_model.keras')
print("Model saved as 'vit_kurdish_handwriting_model.keras'")

# Evaluate the model
test_loss, test_accuracy = vit_model.evaluate(X_test, y_test, verbose=0)
print(f'Test Accuracy: {test_accuracy:.4f}, Test Loss: {test_loss:.4f}')

# Save predictions to a Word document
from docx import Document

predictions = vit_model.predict(X_test)
predicted_classes = np.argmax(predictions, axis=1)
true_classes = np.argmax(y_test, axis=1)

doc = Document()
doc.add_heading('Kurdish Handwriting Recognition - Test Results', level=1)

for i in range(len(predicted_classes)):
    predicted_label = label_encoder.inverse_transform([predicted_classes[i]])[0]
    true_label = label_encoder.inverse_transform([true_classes[i]])[0]
    doc.add_paragraph(f"Sample {i + 1}:")
    doc.add_paragraph(f"Predicted: {predicted_label}")
    doc.add_paragraph(f"Actual: {true_label}")
    doc.add_paragraph("\n")

doc.save('ViT_Test_Results.docx')
print("Results saved to 'ViT_Test_Results.docx'")
