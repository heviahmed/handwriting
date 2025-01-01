import os
import cv2
import numpy as np
from sklearn.model_selection import train_test_split
import pandas as pd
from PIL import Image

from sklearn.preprocessing import LabelEncoder
from tensorflow.keras.utils import to_categorical
from tensorflow.keras.applications import MobileNetV2
from tensorflow.keras.models import Model
from tensorflow.keras.layers import Dense, Dropout, GlobalAveragePooling2D, Input
from tensorflow.keras.optimizers import Adam
from tensorflow.keras.callbacks import ReduceLROnPlateau
from docx import Document

# Paths to your dataset
dataset_path = r'E:\Papers\My Dataset\Images-20241219T162629Z-001\Images'
csv_path = r'E:\Papers\My Dataset\Labels-20241219T162736Z-001\Labels\lines-labels.csv'



#the code for check how many class folders correspond to the IDs, how many classes do not match any folders, and how many IDs don't match any folders.
import os
import pandas as pd
from PIL import Image

# Paths to your dataset
dataset_path = r'E:\Papers\My Dataset\Images-20241219T162629Z-001\Images'
csv_path = r'E:\Papers\My Dataset\Labels-20241219T162736Z-001\Labels\lines-labels.csv'

# Load CSV
csv_data = pd.read_csv(csv_path)

# Normalize CSV IDs
csv_ids = set(csv_data['ID'].str.strip().str.upper())

# List all folders and normalize names
folders = [f for f in os.listdir(dataset_path) if os.path.isdir(os.path.join(dataset_path, f))]
folder_ids = set(f.strip().upper() for f in folders)

# Compute matching IDs
matching_ids = folder_ids & csv_ids

# Debug differences
unmatched_folders = folder_ids - csv_ids
unmatched_csv_ids = csv_ids - folder_ids

print(f"Number of matching IDs (valid IDs): {len(matching_ids)}")
#print(f"Matching IDs: {matching_ids}")
print(f"Number of folders without matching IDs in CSV: {len(unmatched_folders)}")
print(f"Folders without matching IDs: {unmatched_folders}")
print(f"Number of IDs in CSV without matching folders: {len(unmatched_csv_ids)}")
#print(f"Unmatched CSV IDs: {unmatched_csv_ids}")

# Validate images in folders
total_images = 0
invalid_images = []

for folder in folders:
    folder_cleaned = folder.strip().upper()
    if folder_cleaned in matching_ids:  # Only process matching folders
        folder_path = os.path.join(dataset_path, folder)
        images = os.listdir(folder_path)
        for image in images:
            image_path = os.path.join(folder_path, image)
            try:
                with Image.open(image_path) as img:
                    img.verify()  # Check if the image is valid
                total_images += 1
            except Exception as e:
                invalid_images.append(image_path)

print(f"Total valid images loaded: {total_images}")
print(f"Invalid or failed-to-load images: {len(invalid_images)}")
if invalid_images:
    print(f"Invalid images: {invalid_images}")

#the code for checking ids and folders ends here.




# Image size for MobileNetV2
IMG_SIZE = 128


# Load and preprocess the dataset
def load_and_preprocess_dataset():
    """
    Load images and labels from the dataset path and preprocess them for MobileNetV2.
    Returns:
        - images: Preprocessed images (numpy array, shape: (num_samples, 128, 128, 3)).
        - labels: Corresponding labels as integers (numpy array, shape: (num_samples,)).
    """
    import pandas as pd
    from sklearn.preprocessing import LabelEncoder

    # Load labels from the CSV
    labels_df = pd.read_csv(csv_path)

    # Get valid IDs that match folders in dataset_path
    valid_sentence_ids = [folder_name for folder_name in os.listdir(dataset_path) if
                          os.path.isdir(os.path.join(dataset_path, folder_name))]
    filtered_labels_df = labels_df[labels_df['ID'].isin(valid_sentence_ids)].copy()

    # Encode labels
    label_encoder = LabelEncoder()
    filtered_labels_df['label'] = label_encoder.fit_transform(filtered_labels_df['sentence'])

    # Load and preprocess images
    images = []
    labels = []

    for sentence_id in os.listdir(dataset_path):
        sentence_folder = os.path.join(dataset_path, sentence_id)

        if os.path.isdir(sentence_folder) and sentence_id in filtered_labels_df['ID'].values:
            label = filtered_labels_df[filtered_labels_df['ID'] == sentence_id]['label'].values[0]
            for img_file in os.listdir(sentence_folder):
                img_path = os.path.join(sentence_folder, img_file)
                img = cv2.imread(img_path)  # Read the image in RGB
                if img is not None:
                    img = cv2.resize(img, (IMG_SIZE, IMG_SIZE))  # Resize to 128x128
                    img = img / 255.0  # Normalize to [0, 1]
                    images.append(img)
                    labels.append(label)
                else:
                    print(f"Failed to load image: {img_path}")

    images = np.array(images, dtype=np.float32)

    labels = np.array(labels)
    return images, labels, label_encoder


# Load dataset
images, labels, label_encoder = load_and_preprocess_dataset()
print(f"Dataset loaded: {len(images)} samples")
print(f"Image shape: {images[0].shape}, Label shape: {labels.shape}")

# Train-test split
X_train, X_temp, y_train, y_temp = train_test_split(images, labels, test_size=0.3, random_state=42)
X_val, X_test, y_val, y_test = train_test_split(X_temp, y_temp, test_size=0.5, random_state=42)

# Convert labels to one-hot encoding
num_classes = len(np.unique(labels))
y_train = to_categorical(y_train, num_classes)
y_val = to_categorical(y_val, num_classes)
y_test = to_categorical(y_test, num_classes)

# Define the pre-trained model (MobileNetV2)
input_shape = (IMG_SIZE, IMG_SIZE, 3)  # MobileNetV2 expects RGB images
base_model = MobileNetV2(weights='imagenet', include_top=False, input_shape=input_shape)

# Freeze the base model to use it as a feature extractor
base_model.trainable = False

# Add custom layers on top
inputs = Input(shape=input_shape)
x = base_model(inputs, training=False)
x = GlobalAveragePooling2D()(x)  # Pooling layer
x = Dropout(0.5)(x)  # Regularization
x = Dense(256, activation='relu')(x)  # Fully connected layer
outputs = Dense(num_classes, activation='softmax')(x)  # Output layer

# Compile the model
model = Model(inputs, outputs)
model.compile(optimizer=Adam(learning_rate=0.001), loss='categorical_crossentropy', metrics=['accuracy'])

# Summary of the model
model.summary()

# Add a learning rate scheduler
lr_scheduler = ReduceLROnPlateau(monitor='val_loss', factor=0.5, patience=5, min_lr=1e-6)

# Train the model
history = model.fit(
    X_train, y_train,
    epochs=100,
    batch_size=32,
    validation_data=(X_val, y_val),
    callbacks=[lr_scheduler]
)

# Save the trained model
model.save('kurdish_handwriting_mobilenet.keras')
print("Model saved as 'kurdish_handwriting_mobilenet.keras'")

# Evaluate the model on the test set
test_loss, test_accuracy = model.evaluate(X_test, y_test, verbose=0)
print(f'Test Accuracy: {test_accuracy:.4f}, Test Loss: {test_loss:.4f}')

# Save results to a Word document
# Make predictions on the test set
predictions = model.predict(X_test)
predicted_classes = np.argmax(predictions, axis=1)
true_classes = np.argmax(y_test, axis=1)









# Create a Word document
doc = Document()
doc.add_heading('Kurdish Handwriting Recognition - Test Results', level=1)

for i in range(len(predicted_classes)):
    predicted_label = label_encoder.inverse_transform([predicted_classes[i]])[0]
    true_label = label_encoder.inverse_transform([true_classes[i]])[0]

    doc.add_paragraph(f"Sample {i + 1}:")
    doc.add_paragraph(f"Predicted: {predicted_label}")
    doc.add_paragraph(f"Actual: {true_label}")
    doc.add_paragraph("\n")

# Save the document
doc.save('Test_Results.docx')
#print("Results saved to 'Test_Results.docx'")
