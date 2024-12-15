import os
import numpy as np
from tensorflow.keras.applications import InceptionV3
from tensorflow.keras.models import Model
from tensorflow.keras.layers import Dense, Dropout, GlobalAveragePooling2D, Input
from tensorflow.keras.optimizers import Adam
from tensorflow.keras.callbacks import ReduceLROnPlateau
from tensorflow.keras.utils import to_categorical
from sklearn.preprocessing import LabelEncoder
from sklearn.model_selection import train_test_split
import pandas as pd
import cv2

# Paths to your dataset
dataset_path = r'C:\Users\BEST TECH\Downloads\SUE+UKH'
csv_path = r'C:\Users\BEST TECH\Downloads\UKH+SUE-labels.csv'

# Image size for InceptionV3
IMG_SIZE = 150  # Smaller size to reduce memory usage

# Load and preprocess dataset
def load_and_preprocess_dataset():
    labels_df = pd.read_csv(csv_path)

    # Get valid folder IDs
    valid_sentence_ids = [folder_name for folder_name in os.listdir(dataset_path) if
                          os.path.isdir(os.path.join(dataset_path, folder_name))]
    filtered_labels_df = labels_df[labels_df['ID'].isin(valid_sentence_ids)].copy()

    # Encode labels
    label_encoder = LabelEncoder()
    filtered_labels_df['label'] = label_encoder.fit_transform(filtered_labels_df['sentence'])

    images = []
    labels = []
    valid_extensions = ('.jpg', '.jpeg', '.png', '.bmp')  # Valid image file extensions

    for sentence_id in os.listdir(dataset_path):
        sentence_folder = os.path.join(dataset_path, sentence_id)
        if os.path.isdir(sentence_folder) and sentence_id in filtered_labels_df['ID'].values:
            label = filtered_labels_df[filtered_labels_df['ID'] == sentence_id]['label'].values[0]
            for img_file in os.listdir(sentence_folder):
                img_path = os.path.join(sentence_folder, img_file)

                # Check for valid image file extension
                if not img_file.lower().endswith(valid_extensions):
                    print(f"Skipping non-image file: {img_path}")
                    continue

                # Attempt to load the image
                try:
                    img = cv2.imread(img_path)  # Read image as RGB
                    if img is None:
                        raise ValueError(f"Cannot read image: {img_path}")
                    img = cv2.resize(img, (IMG_SIZE, IMG_SIZE)) / 255.0
                    images.append(img)
                    labels.append(label)
                except Exception as e:
                    print(f"Skipping corrupted or unsupported image: {img_path}, Error: {e}")

    images = np.array(images, dtype=np.float32)
    labels = np.array(labels)
    print(f"Loaded {len(images)} images and {len(labels)} labels.")
    return images, labels, label_encoder

# Load dataset
images, labels, label_encoder = load_and_preprocess_dataset()

# Train-test split
X_train, X_temp, y_train, y_temp = train_test_split(images, labels, test_size=0.3, random_state=42)
X_val, X_test, y_val, y_test = train_test_split(X_temp, y_temp, test_size=0.5, random_state=42)

# Convert labels to categorical
num_classes = len(np.unique(labels))
y_train = to_categorical(y_train, num_classes)
y_val = to_categorical(y_val, num_classes)
y_test = to_categorical(y_test, num_classes)

# Define the InceptionV3 Model
input_shape = (IMG_SIZE, IMG_SIZE, 3)  # InceptionV3 expects RGB images
base_model = InceptionV3(weights='imagenet', include_top=False, input_shape=input_shape)

# Freeze the base model
base_model.trainable = False

# Add custom layers on top
inputs = Input(shape=input_shape)
x = base_model(inputs, training=False)
x = GlobalAveragePooling2D()(x)
x = Dropout(0.5)(x)
x = Dense(256, activation='relu')(x)
outputs = Dense(num_classes, activation='softmax')(x)

# Compile the model
model = Model(inputs, outputs)
model.compile(optimizer=Adam(learning_rate=0.001), loss='categorical_crossentropy', metrics=['accuracy'])

# Print model summary
model.summary()

# Add learning rate scheduler
lr_scheduler = ReduceLROnPlateau(monitor='val_loss', factor=0.5, patience=5, min_lr=1e-6)

# Train the model
history = model.fit(
    X_train, y_train,
    epochs=20,
    batch_size=32,
    validation_data=(X_val, y_val),
    callbacks=[lr_scheduler]
)

# Save the trained model
model.save('inceptionv3_kurdish_handwriting.keras')
print("Model saved as 'inceptionv3_kurdish_handwriting.keras'")

# Evaluate the model on the test set
test_loss, test_accuracy = model.evaluate(X_test, y_test, verbose=0)
print(f'Test Accuracy: {test_accuracy:.4f}, Test Loss: {test_loss:.4f}')

# Save predictions to a Word document
from docx import Document

# Get predictions
predictions = model.predict(X_test, verbose=0)
predicted_classes = np.argmax(predictions, axis=1)

# Get true classes
true_classes = np.argmax(y_test, axis=1)

# Create a Word document
doc = Document()
doc.add_heading('Kurdish Handwriting Recognition - Test Results', level=1)

for i in range(len(predicted_classes)):
    predicted_label = label_encoder.inverse_transform([predicted_classes[i]])[0]
    true_label = label_encoder.inverse_transform([true_classes[i]])[0]
    doc.add_paragraph(f"Sample {i + 1}:")
    doc.add_paragraph(f"Predicted: {predicted_label}")
    doc.add_paragraph(f"Actual: {true_label}")
    doc.add_paragraph("\n")

# Save the document
doc.save('InceptionV3_Test_Results.docx')
print("Results saved to 'InceptionV3_Test_Results.docx'")
